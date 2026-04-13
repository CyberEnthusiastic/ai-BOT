"""Document tool — read/write .docx (python-docx), .xlsx (openpyxl), read .pdf (PyPDF2).

No mock mode needed — these are purely local file operations.
All methods are async; file I/O runs in an executor.
"""

from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any


# ── DocumentTool ──────────────────────────────────────────────────────────────
class DocumentTool:

    # ── read_docx ─────────────────────────────────────────────────────────────
    async def read_docx(self, path: str) -> str:
        """Read a Word document (.docx) and return its text content.

        Args:
            path: File path to the .docx file.

        Returns:
            Full text of the document, paragraphs separated by newlines.
        """
        def _read() -> str:
            try:
                from docx import Document  # type: ignore
            except ImportError as exc:
                raise RuntimeError(
                    "python-docx not installed. Run: pip install python-docx"
                ) from exc

            p = Path(path)
            if not p.exists():
                return f"File not found: {path}"
            if p.suffix.lower() != ".docx":
                return f"Not a .docx file: {path}"

            doc = Document(str(p))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Also extract text from tables
            table_rows: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))

            text = "\n".join(paragraphs)
            if table_rows:
                text += "\n\n[Tables]\n" + "\n".join(table_rows)
            return text or "(Document is empty)"

        return await asyncio.get_event_loop().run_in_executor(None, _read)

    # ── write_docx ────────────────────────────────────────────────────────────
    async def write_docx(
        self,
        path: str,
        content: str,
        title: str = "",
        overwrite: bool = False,
    ) -> str:
        """Write or create a Word document (.docx).

        Content is split by blank lines into paragraphs. Lines starting with
        '# ' become headings, '## ' become subheadings.

        Args:
            path: Destination file path (must end in .docx).
            content: Text content. Use blank lines to separate paragraphs.
                     Lines starting with '# ' become Heading 1,
                     '## ' become Heading 2.
            title: Optional document title added at the top.
            overwrite: If False (default), raises error if file already exists.

        Returns:
            Confirmation with the saved path.
        """
        def _write() -> str:
            try:
                from docx import Document  # type: ignore
                from docx.shared import Pt  # type: ignore
            except ImportError as exc:
                raise RuntimeError(
                    "python-docx not installed. Run: pip install python-docx"
                ) from exc

            p = Path(path)
            if not overwrite and p.exists():
                return f"File already exists: {path}. Pass overwrite=True to replace."
            if p.suffix.lower() != ".docx":
                return f"Path must end in .docx: {path}"

            p.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()

            if title:
                doc.add_heading(title, level=0)

            for line in content.split("\n"):
                stripped = line.rstrip()
                if stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(stripped)

            doc.save(str(p))
            return f"Document saved: {path} ({p.stat().st_size} bytes)"

        return await asyncio.get_event_loop().run_in_executor(None, _write)

    # ── read_xlsx ─────────────────────────────────────────────────────────────
    async def read_xlsx(self, path: str, sheet_name: str = "") -> str:
        """Read an Excel spreadsheet (.xlsx) and return its data as text.

        Args:
            path: File path to the .xlsx file.
            sheet_name: Sheet to read (default: first sheet).

        Returns:
            Tab-separated rows with a header, one row per line.
        """
        def _read() -> str:
            try:
                import openpyxl  # type: ignore
            except ImportError as exc:
                raise RuntimeError(
                    "openpyxl not installed. Run: pip install openpyxl"
                ) from exc

            p = Path(path)
            if not p.exists():
                return f"File not found: {path}"
            if p.suffix.lower() not in (".xlsx", ".xlsm"):
                return f"Not an Excel file (.xlsx/.xlsm): {path}"

            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
            if ws is None:
                return f"Sheet '{sheet_name}' not found. Available: {wb.sheetnames}"

            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                rows.append("\t".join(cells))

            wb.close()
            if not rows:
                return "(Sheet is empty)"

            sheet_label = ws.title if ws.title else "Sheet1"
            return f"Sheet: {sheet_label}\n\n" + "\n".join(rows)

        return await asyncio.get_event_loop().run_in_executor(None, _read)

    # ── write_xlsx ────────────────────────────────────────────────────────────
    async def write_xlsx(
        self,
        path: str,
        rows: list[list[Any]],
        sheet_name: str = "Sheet1",
        headers: list[str] | None = None,
        overwrite: bool = False,
    ) -> str:
        """Write data to an Excel spreadsheet (.xlsx).

        Args:
            path: Destination file path (must end in .xlsx).
            rows: List of rows, each row is a list of cell values.
            sheet_name: Name for the worksheet.
            headers: Optional list of column header strings (added as first row).
            overwrite: If False (default), raises error if file already exists.

        Returns:
            Confirmation with the saved path and row count.
        """
        def _write() -> str:
            try:
                import openpyxl  # type: ignore
                from openpyxl.styles import Font  # type: ignore
            except ImportError as exc:
                raise RuntimeError(
                    "openpyxl not installed. Run: pip install openpyxl"
                ) from exc

            p = Path(path)
            if not overwrite and p.exists():
                return f"File already exists: {path}. Pass overwrite=True to replace."
            if p.suffix.lower() not in (".xlsx", ".xlsm"):
                return f"Path must end in .xlsx: {path}"

            p.parent.mkdir(parents=True, exist_ok=True)
            wb = openpyxl.Workbook()
            ws = wb.active
            if ws is None:
                ws = wb.create_sheet()
            ws.title = sheet_name

            if headers:
                ws.append(headers)
                # Bold the header row
                for cell in ws[1]:
                    cell.font = Font(bold=True)

            for row in rows:
                ws.append(row)

            wb.save(str(p))
            row_count = len(rows) + (1 if headers else 0)
            return f"Spreadsheet saved: {path} ({row_count} rows, {p.stat().st_size} bytes)"

        return await asyncio.get_event_loop().run_in_executor(None, _write)

    # ── read_pdf ──────────────────────────────────────────────────────────────
    async def read_pdf(self, path: str, max_pages: int = 50) -> str:
        """Read a PDF file and extract its text content.

        Args:
            path: File path to the .pdf file.
            max_pages: Maximum number of pages to read (default 50).

        Returns:
            Extracted text, page-by-page with page markers.
        """
        def _read() -> str:
            try:
                import PyPDF2  # type: ignore
            except ImportError as exc:
                raise RuntimeError(
                    "PyPDF2 not installed. Run: pip install PyPDF2"
                ) from exc

            p = Path(path)
            if not p.exists():
                return f"File not found: {path}"
            if p.suffix.lower() != ".pdf":
                return f"Not a PDF file: {path}"

            pages: list[str] = []
            with open(str(p), "rb") as f:
                reader = PyPDF2.PdfReader(f)
                total = len(reader.pages)
                limit = min(total, max_pages)
                for i in range(limit):
                    text = reader.pages[i].extract_text() or ""
                    pages.append(f"[Page {i + 1}]\n{text.strip()}")

            suffix = f"\n\n(Showing {limit} of {total} pages)" if total > max_pages else ""
            return "\n\n".join(pages) + suffix if pages else "(No text extracted from PDF)"

        return await asyncio.get_event_loop().run_in_executor(None, _read)
