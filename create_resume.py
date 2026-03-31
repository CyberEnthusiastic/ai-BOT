"""
Generate Mohith Vasamsetti's elite master resume as .docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(9.5)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Helper: add horizontal line
def add_hr(doc, color='1F4E79'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# Helper: add section heading
def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue
    run.font.name = 'Calibri'
    add_hr(doc)

# Helper: add job header
def add_job_header(doc, company, title, location, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)

    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    run = p.add_run('  |  ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    run = p.add_run(title)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    run = p.add_run('  |  ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    run = p.add_run(location)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    run = p.add_run('  |  ')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    run = p.add_run(dates)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

# Helper: add bullet point
def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)

    # Split on bold markers (em dash separates action from result)
    if '\u2014' in text:
        parts = text.split('\u2014', 1)
        run = p.add_run('\u2022 ')
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'

        # Bold the action part
        run = p.add_run(parts[0].strip())
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'

        run = p.add_run(' \u2014 ')
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'

        run = p.add_run(parts[1].strip())
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
    else:
        run = p.add_run('\u2022 ' + text)
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'

# Helper: add skills line
def add_skill_line(doc, category, skills):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0.5)

    run = p.add_run(category + ': ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'

    run = p.add_run(skills)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'


# ==================== BUILD RESUME ====================

# --- NAME ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Mohith Vasamsetti')
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(31, 78, 121)

# --- CONTACT ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('+1 (254) 284-4847  |  vmna56@gmail.com  |  United States  |  Open to Relocate')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(80, 80, 80)

# --- SUMMARY ---
add_section_heading(doc, 'SUMMARY')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)

summary_text = (
    'Security Engineer with 5+ years securing enterprise-scale cloud infrastructure, '
    'application security, and SOC/SIEM operations across Fortune 500, FinTech, and critical '
    'infrastructure organizations. Expert in AWS, Kubernetes, Splunk, Terraform, and Python-driven '
    'security automation at scale. Proven track record building DevSecOps pipelines, SOAR playbooks, '
    'threat detection programs, and vulnerability management workflows under NIST CSF, PCI-DSS, SOC 2, '
    'and ISO 27001 frameworks. Deep expertise in IAM governance, container security, zero-trust architecture, '
    'privacy engineering, and incident response for globally distributed engineering teams.'
)

# Bold key terms
bold_terms = [
    'Security Engineer', 'cloud infrastructure', 'application security', 'SOC/SIEM operations',
    'AWS', 'Kubernetes', 'Splunk', 'Terraform', 'Python', 'DevSecOps pipelines', 'SOAR playbooks',
    'threat detection', 'vulnerability management', 'NIST CSF', 'PCI-DSS', 'SOC 2', 'ISO 27001',
    'IAM governance', 'container security', 'zero-trust architecture', 'incident response'
]

run = p.add_run(summary_text)
run.font.size = Pt(9.5)
run.font.name = 'Calibri'

# --- WORK EXPERIENCE ---
add_section_heading(doc, 'WORK EXPERIENCE')

# Costco
add_job_header(doc, 'Costco Wholesale', 'Security Engineering Consultant', 'Seattle, WA', 'Nov 2025 \u2013 Present')
costco_bullets = [
    'Architected enterprise Splunk SIEM with 40+ correlation rules and SOAR playbooks \u2014 automating 80% of Tier-1 alert triage and reducing mean time to detect (MTTD) by 52% across 10,000+ endpoints',
    'Deployed AWS Security Hub, GuardDuty, WAF, and Macie across 200+ production accounts \u2014 blocking 1M+ malicious requests monthly and achieving CIS Benchmark compliance with 61% fewer critical misconfigurations',
    'Built Terraform IaC security modules and policy-as-code guardrails \u2014 enforcing encryption, IAM least privilege, and network segmentation with 100% compliance across all provisioned cloud infrastructure',
    'Integrated GitHub Enterprise SAST, SCA, and secrets detection using Snyk, Semgrep, and GitGuardian across 150+ repositories \u2014 remediating 94% of critical findings within SLA and blocking 100% of credential exposure',
    'Automated Python-driven vulnerability scanning, triage, and compliance reporting workflows \u2014 reducing critical CVE remediation from 30 days to under 7 days and saving 25+ hours per quarterly audit cycle',
]
for b in costco_bullets:
    add_bullet(doc, b)

# Amazon
add_job_header(doc, 'Amazon', 'Security Engineer, Cloud & Detection Engineering', 'Arlington, VA', 'Jan 2025 \u2013 Oct 2025')
amazon_bullets = [
    'Engineered AWS-native detection pipelines using CloudTrail, Config, GuardDuty, and Lambda \u2014 processing 1B+ events daily with 99.7% anomalous activity detection across multi-region production environments',
    'Secured Kubernetes workloads at scale \u2014 implementing Pod Security Standards, OPA Gatekeeper, Falco runtime monitoring, and RBAC hardening across 40+ production microservices',
    'Automated CI/CD security pipeline stages using GitHub Actions \u2014 embedding SAST, DAST, container scanning, and dependency review that blocked 72% more vulnerabilities pre-deployment with zero-downtime releases',
    'Built Python/Boto3 security automation \u2014 including IAM least-privilege analysis eliminating 340+ over-privileged roles, SOAR incident enrichment reducing analyst triage time by 60%, and threat hunting scripts identifying 3 unknown persistence mechanisms',
    'Designed Prometheus and Grafana security observability dashboards \u2014 tracking MTTD, MTTR, vulnerability SLAs, and compliance posture in real time for 12 cross-functional engineering teams',
]
for b in amazon_bullets:
    add_bullet(doc, b)

# Axis Bank
add_job_header(doc, 'Axis Bank', 'Security Engineer, SOC & Enterprise Security', 'India', 'Aug 2020 \u2013 Jul 2023')
axis_bullets = [
    'Deployed Splunk SIEM ingesting 50M+ daily events \u2014 with custom correlation rules for financial fraud, insider threat, privilege escalation, and credential abuse across 15 critical banking applications',
    'Enforced PCI-DSS Level 1 and ISO 27001 compliance programs \u2014 achieving zero audit exceptions across 3 consecutive annual QSA reviews and external penetration tests',
    'Conducted threat hunting exercises using MITRE ATT&CK framework \u2014 identifying 8 previously undetected persistence mechanisms in banking application servers',
]
for b in axis_bullets:
    add_bullet(doc, b)

# --- SKILLS ---
add_section_heading(doc, 'SKILLS')
add_skill_line(doc, 'Cloud & Infrastructure', 'AWS (GuardDuty, Security Hub, WAF, Macie, CloudTrail, IAM, Lambda, KMS, Config, VPC, S3), Azure, GCP, Kubernetes, Docker, Terraform, Helm')
add_skill_line(doc, 'SIEM & Detection', 'Splunk Enterprise Security, SOAR, Elastic SIEM, Falco, OPA Gatekeeper, Prometheus, Grafana, CrowdStrike, SentinelOne')
add_skill_line(doc, 'AppSec & DevSecOps', 'SAST, DAST (OWASP ZAP, Burp Suite), SCA, GitHub Enterprise, CI/CD Security Gates, IaC Scanning, Secrets Management')
add_skill_line(doc, 'IAM & Zero Trust', 'Okta, OAuth2/OIDC/SAML, AWS IAM/SCPs, RBAC, PAM, MFA, Identity Lifecycle Management, Endpoint Security')
add_skill_line(doc, 'Scripting & Automation', 'Python, Go, Bash, PowerShell, Rust, SQL, REST APIs, Boto3')
add_skill_line(doc, 'Frameworks & Compliance', 'NIST CSF, NIST 800-53, PCI-DSS, SOC 2, ISO 27001, HIPAA, GDPR/CCPA, NERC CIP, CIS Benchmarks, MITRE ATT&CK, OWASP Top 10')
add_skill_line(doc, 'Relevant Skills', 'Incident Response, Threat Hunting, Vulnerability Management, Privacy Engineering, DLP, CSPM, Penetration Testing, Bug Bounty, AI/ML Security')

# --- EDUCATION ---
add_section_heading(doc, 'EDUCATION')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(1)
run = p.add_run('M.S., Cybersecurity & Information Assurance')
run.bold = True
run.font.size = Pt(9.5)
run.font.name = 'Calibri'
run = p.add_run('  |  University of Central Missouri, USA')
run.font.size = Pt(9.5)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(1)
run = p.add_run('B.Tech., Computer Science')
run.bold = True
run.font.size = Pt(9.5)
run.font.name = 'Calibri'
run = p.add_run('  |  DVR & DrHS MIC College of Technology, India')
run.font.size = Pt(9.5)
run.font.name = 'Calibri'

# --- CERTIFICATIONS ---
add_section_heading(doc, 'CERTIFICATIONS')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('AWS Certified Security - Specialty  |  CompTIA Security+  |  CISA')
run.font.size = Pt(9.5)
run.font.name = 'Calibri'

# Save
output_path = os.path.join('all resumes', 'default', 'resume.docx')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Resume saved to: {output_path}")
